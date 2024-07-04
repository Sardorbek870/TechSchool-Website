import random

from docx import Document
option1 = []
option2 = []
option3 = []


def read_word_file(file_path):
    doc = Document(file_path)

    for paragraph in doc.paragraphs:
        lines = paragraph.text.split("\n")
        for line in lines:
            if line.startswith("1"):
                option1.append(line)
            elif line.startswith("2"):
                option2.append(line)
            elif line.startswith("3"):
                option3.append(line)


read_word_file("variant1.docx")


def write_to_word_file(file_path, lines):
    doc = Document()

    for line in lines:
        doc.add_paragraph(line)
        doc.add_paragraph("\n")

    doc.save(file_path)

file_path = r"C:\Users\smuro\OneDrive\Desktop\savollar.docx"
doc = Document()
for i in range(1, 40):
    lines = []
    lines.append(f"{" " * 15}Variant {i}")
    lines.append(random.choice(option1))
    lines.append(random.choice(option2))
    lines.append(random.choice(option3))

    for line in lines:
        doc.add_paragraph(line)


doc.save(file_path)
print("Done Successfully")




